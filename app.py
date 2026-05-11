import streamlit as st
import pandas as pd
import pdfplumber
import docx
from docx.shared import Pt, RGBColor
import re
import io
from datetime import datetime

st.set_page_config(page_title="Auditor JPV - Revisión Individual", layout="wide")

# --- ESTADO DE SESIÓN (Persistencia de la Base Maestra) ---
if 'df_maestro' not in st.session_state:
    st.session_state.df_maestro = None

st.title("🔎 Auditor Individual de Informes")
st.markdown("Revisión uno a uno contra el Reporte de Acciones persistente.")

# --- 1. FUNCIONES DE EXTRACCIÓN ---
def extraer_texto_pdf(archivo):
    texto = ""
    try:
        with pdfplumber.open(archivo) as pdf:
            for pagina in pdf.pages:
                txt = pagina.extract_text(layout=True)
                if txt: texto += txt + "\n"
    except Exception as e:
        texto = f"Error al leer PDF: {e}"
    return texto

def extraer_texto_docx(archivo):
    texto = ""
    try:
        doc = docx.Document(archivo)
        for para in doc.paragraphs:
            texto += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                texto += " | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]) + "\n"
    except Exception as e:
        texto = f"Error al leer DOCX: {e}"
    return texto

def limpiar_monto(texto_monto):
    if not texto_monto: return 0.0
    limpio = re.sub(r'[^\d,\.-]', '', str(texto_monto))
    limpio = limpio.replace('.', '').replace(',', '.')
    try:
        return float(limpio)
    except:
        return 0.0

def extraer_datos_informe(texto):
    datos = {
        "Liquidacion": None, "Poliza": None, "Compania": None, "Ajustador": None,
        "Fecha_Siniestro": None, "Fecha_Denuncia": None, "Fecha_Inspeccion": None,
        "Divisa": None, "Perdida_Bruta": 0.0
    }
    
    # Llave y Forma
    match_liq = re.search(r'(?:LIQUIDACI[OÓ]N Nº|Ref\. JPV\s*:)\s*(\d+)', texto, re.IGNORECASE)
    if match_liq: datos["Liquidacion"] = match_liq.group(1).strip()
        
    match_pol = re.search(r'(?:Nº Póliza|Póliza Nº|Póliza número)[\s:]*([A-Za-z0-9-]+)', texto, re.IGNORECASE)
    if match_pol: datos["Poliza"] = match_pol.group(1).strip()
    
    match_comp = re.search(r'(?:ASEGURADOR|Compañía)[\s:]*([^\n|]+)', texto, re.IGNORECASE)
    if match_comp: datos["Compania"] = match_comp.group(1).strip()

    # Ajustador
    match_ajust = re.search(r'Ajustador a cargo[\s:]*([^\n|]+)', texto, re.IGNORECASE)
    if match_ajust: datos["Ajustador"] = match_ajust.group(1).strip()

    # Fechas
    match_fsin = re.search(r'Fecha de Siniestro[\s:]*([\d\-]+)', texto, re.IGNORECASE)
    if match_fsin: datos["Fecha_Siniestro"] = match_fsin.group(1).strip()
        
    match_fden = re.search(r'Fecha Denuncia[\s:]*([\d\-]+)', texto, re.IGNORECASE)
    if match_fden: datos["Fecha_Denuncia"] = match_fden.group(1).strip()

    match_fins = re.search(r'Fecha Inspección[\s:]*([\d\-]+)', texto, re.IGNORECASE)
    if match_fins: datos["Fecha_Inspeccion"] = match_fins.group(1).strip()

    # Monto Bruto (Pérdida Bruta, Reserva Determinada, Pérdida Estimada, Pérdida Probable BF)
    matches_bruta = re.findall(r'(?:Pérdida Bruta|Reserva Determinada|Pérdida Estimada|Pérdida Probable BF)[^\d]*([\d\.,]+)', texto, re.IGNORECASE)
    if matches_bruta: datos["Perdida_Bruta"] = limpiar_monto(matches_bruta[-1])
    
    # Divisa
    match_div = re.search(r'(UF|US\$|USD|\$)', texto, re.IGNORECASE)
    if match_div: datos["Divisa"] = match_div.group(1).upper()
        
    return datos

# --- 2. GESTIÓN DE LA BASE MAESTRA EN SIDEBAR ---
with st.sidebar:
    st.header("⚙️ Configuración")
    archivo_excel = st.file_uploader("Actualizar Reporte de Acciones (Excel)", type=["xlsx"])
    
    if archivo_excel:
        for i in range(10):
            try:
                df_temp = pd.read_excel(archivo_excel, skiprows=i)
                df_temp.columns = [str(c).strip() for c in df_temp.columns]
                if 'Número de caso' in df_temp.columns:
                    df_temp['Número de caso'] = df_temp['Número de caso'].astype(str).str.strip().str.replace(r'\.0$', '', regex=True)
                    st.session_state.df_maestro = df_temp.dropna(how='all', axis=0)
                    st.success("✅ Base Maestra cargada y guardada en sesión.")
                    break
            except: continue

    if st.session_state.df_maestro is not None:
        st.write(f"📊 **Base actual:** {len(st.session_state.df_maestro)} casos registrados.")
        if st.button("Limpiar Base de Datos"):
            st.session_state.df_maestro = None
            st.rerun()

# --- 3. GENERADOR DE CERTIFICADO INDIVIDUAL (WORD) ---
def generar_word_individual(resultado):
    doc = docx.Document()
    
    doc.add_heading('CERTIFICADO DE AUDITORÍA - REVISIÓN INDIVIDUAL', 0)
    doc.add_paragraph(f"Fecha de Emisión: {datetime.now().strftime('%d-%m-%Y %H:%M')}\nAuditoría Realizada por: Módulo de Control Automático JPV")
    
    doc.add_heading('1. DATOS DEL CASO AUDITADO', level=1)
    doc.add_paragraph(f"Liquidación Nº (Caso): {resultado['N° Caso']}\nDocumento Revisado: {resultado['Documento']}\nAjustador Senior Asignado: {resultado['Ajustador_Sistema']}", style='List Bullet')

    estado_gral = "❌ CON OBSERVACIONES" if "ERROR" in resultado['Validación Forma'] or "ERROR" in resultado['Validación Fechas'] or "DESVIACIÓN" in resultado['Desviación Reserva'] else "✅ APROBADO"
    doc.add_heading(f'2. RESULTADOS DE LA AUDITORÍA: {estado_gral}', level=1)

    # A. IDENTIFICACIÓN Y FORMA
    doc.add_heading('A. IDENTIFICACIÓN Y FORMA', level=2)
    doc.add_paragraph(f"Compañía de Seguros: {resultado['Compania']}", style='List Bullet')
    doc.add_paragraph(f"Ajustador Firmante: {resultado['Ajustador']}", style='List Bullet')
    p_poliza = doc.add_paragraph(style='List Bullet')
    p_poliza.add_run(f"Póliza de Seguros: {resultado['Validación Forma']}")
    if "ERROR" in resultado['Validación Forma']:
        for det in resultado['Detalles_Forma']:
            p_det = doc.add_paragraph(f"  • {det}")
            p_det.runs[0].font.color.rgb = RGBColor(200, 0, 0)

    # B. REVISIÓN CRONOLÓGICA
    doc.add_heading('B. REVISIÓN CRONOLÓGICA', level=2)
    p_fecha = doc.add_paragraph(style='List Bullet')
    p_fecha.add_run(f"Fechas (Ocurrencia/Denuncia/Inspección): {resultado['Validación Fechas']}")
    if "ERROR" in resultado['Validación Fechas']:
        for det in resultado['Detalles_Fechas']:
            p_det = doc.add_paragraph(f"  • {det}")
            p_det.runs[0].font.color.rgb = RGBColor(200, 0, 0)

    # C. REVISIÓN FINANCIERA Y MONEDA
    doc.add_heading('C. REVISIÓN FINANCIERA Y MONEDA', level=2)
    doc.add_paragraph(f"Divisa: {resultado['Divisa']}", style='List Bullet')
    p_bruta = doc.add_paragraph(style='List Bullet')
    p_bruta.add_run(f"Pérdida Bruta: {resultado['Desviación Reserva']}")
    if "DESVIACIÓN" in resultado['Desviación Reserva']:
        for det in resultado['Detalles_Bruta']:
            p_det = doc.add_paragraph(f"  • {det}")
            p_det.runs[0].font.color.rgb = RGBColor(200, 100, 0)

    doc.add_paragraph("\nAcción Requerida: " + ("Se solicita al ajustador la rectificación de los puntos marcados antes de la emisión final." if estado_gral != "✅ APROBADO" else "El informe cumple con los estándares revisados."))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer

# --- 4. FLUJO DE AUDITORÍA PRINCIPAL ---
if st.session_state.df_maestro is not None:
    st.subheader("📄 Sube el Informe a Revisar")
    archivo_informe = st.file_uploader("Sube un único Informe (PDF o DOCX)", type=["pdf", "docx"])

    if archivo_informe:
        texto = extraer_texto_pdf(archivo_informe) if archivo_informe.name.lower().endswith('.pdf') else extraer_texto_docx(archivo_informe)
        datos_doc = extraer_datos_informe(texto)
        
        if not datos_doc["Liquidacion"]:
            st.error("❌ No se pudo detectar el número de liquidación en el documento.")
        else:
            df = st.session_state.df_maestro
            fila = df[df['Número de caso'] == datos_doc["Liquidacion"]]
            
            if fila.empty:
                st.warning(f"⚠️ El caso {datos_doc['Liquidacion']} no se encuentra en el Reporte de Acciones.")
            else:
                fila = fila.iloc[0]
                
                detalles_forma = []
                detalles_fechas = []
                detalles_bruta = []
                
                # --- A. IDENTIFICACIÓN Y FORMA ---
                v_forma = "✅ OK"
                poliza_sist = str(fila.get('Póliza de seguros', '')).strip()
                if poliza_sist.upper() not in str(datos_doc["Poliza"]).upper() and poliza_sist != "nan" and poliza_sist != "":
                    v_forma = "❌ ERROR"
                    detalles_forma.append(f"Documento indica: {datos_doc['Poliza']} | Sistema indica: {poliza_sist}")
                
                comp_sist = str(fila.get('Compañía de seguros', '')).strip()
                v_comp = "✅ OK" if (comp_sist.upper() in str(datos_doc["Compania"]).upper() or not datos_doc["Compania"]) else "❌ ERROR"
                if v_comp == "❌ ERROR": detalles_forma.append(f"Compañía en Doc: {datos_doc['Compania']} | Sist: {comp_sist}")

                ajust_sist = str(fila.get('Ajustador senior', '')).strip()
                v_ajust = "✅ OK" if (ajust_sist.upper() in str(datos_doc["Ajustador"]).upper() or not datos_doc["Ajustador"]) else "⚠️ WARNING"

                # --- B. FECHAS ---
                v_fechas = "✅ OK"
                f_ocurr = str(fila.get('Fecha de ocurrencia', '')).strip()
                if datos_doc["Fecha_Siniestro"] and datos_doc["Fecha_Siniestro"][:10] not in f_ocurr:
                    v_fechas = "❌ ERROR"
                    detalles_fechas.append(f"Ocurrencia: Doc({datos_doc['Fecha_Siniestro'][:10]}) vs Sist({f_ocurr})")

                f_denun = str(fila.get('Fecha de denuncio', '')).strip()
                if datos_doc["Fecha_Denuncia"] and datos_doc["Fecha_Denuncia"][:10] not in f_denun:
                    v_fechas = "❌ ERROR"
                    detalles_fechas.append(f"Denuncio: Doc({datos_doc['Fecha_Denuncia'][:10]}) vs Sist({f_denun})")

                f_insp = str(fila.get('Fecha de inspección', '')).strip()
                if datos_doc["Fecha_Inspeccion"] and datos_doc["Fecha_Inspeccion"][:10] not in f_insp:
                    v_fechas = "❌ ERROR"
                    detalles_fechas.append(f"Inspección: Doc({datos_doc['Fecha_Inspeccion'][:10]}) vs Sist({f_insp})")

                # --- C. FINANCIERA ---
                v_bruta = "✅ OK"
                bruta_sist = limpiar_monto(fila.get('Perdida bruta (en moneda del caso)', 0))
                if abs(datos_doc["Perdida_Bruta"] - bruta_sist) > 1.0 and bruta_sist > 0:
                    v_bruta = "⚠️ DESVIACIÓN DETECTADA"
                    detalles_bruta.append(f"Monto Doc: {datos_doc['Perdida_Bruta']:,.2f} | Monto Sist: {bruta_sist:,.2f}")

                div_sist = str(fila.get('Divisa', '')).strip()
                v_div = "✅ OK" if (div_sist.upper() in str(datos_doc["Divisa"]).upper()) else "❌ ERROR"
                if v_div == "❌ ERROR": detalles_bruta.append(f"Divisa Doc: {datos_doc['Divisa']} | Sist: {div_sist}")

                # --- RESULTADO FINAL ---
                resultado_final = {
                    "Documento": archivo_informe.name,
                    "N° Caso": datos_doc["Liquidacion"],
                    "Ajustador_Sistema": ajust_sist,
                    "Validación Forma": v_forma,
                    "Detalles_Forma": detalles_forma,
                    "Compania": f"{v_comp} ({datos_doc['Compania']})",
                    "Ajustador": f"{v_ajust} ({datos_doc['Ajustador']})",
                    "Validación Fechas": v_fechas,
                    "Detalles_Fechas": detalles_fechas,
                    "Desviación Reserva": v_bruta,
                    "Detalles_Bruta": detalles_bruta,
                    "Divisa": f"{v_div} ({datos_doc['Divisa']})"
                }

                # Pantalla
                estado_visual = "✅ Aprobado" if "ERROR" not in v_forma and "ERROR" not in v_fechas and "DESVIACIÓN" not in v_bruta else "❌ Con Observaciones"
                st.success(f"Auditoría del caso {datos_doc['Liquidacion']} completada. Estado: {estado_visual}")
                
                col1, col2, col3 = st.columns(3)
                col1.metric("Forma y Póliza", "❌ Error" if detalles_forma else "✅ OK")
                col2.metric("Fechas", "❌ Error" if detalles_fechas else "✅ OK")
                col3.metric("Pérdida Bruta", "⚠️ Desviación" if detalles_bruta else "✅ OK")

                if detalles_forma or detalles_fechas or detalles_bruta:
                    st.error("📋 Hallazgos detectados:")
                    for d in detalles_forma + detalles_fechas + detalles_bruta:
                        st.write(f"• {d}")

                # Descarga
                word_buf = generar_word_individual(resultado_final)
                st.download_button(
                    label="📄 Descargar Certificado de Auditoría (Word)",
                    data=word_buf.getvalue(),
                    file_name=f"Auditoria_{datos_doc['Liquidacion']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
else:
    st.info("👈 Por favor, carga primero el Reporte de Acciones en la barra lateral para comenzar.")
